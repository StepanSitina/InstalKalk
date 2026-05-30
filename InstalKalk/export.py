from docx import Document
from docx.shared import Pt, Cm
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill
import os
import datetime

os.makedirs("exports", exist_ok=True)

def generate_word(project_name: str, items: list) -> str:
    doc = Document()
    doc.add_heading(f'Cenová nabídka: {project_name}', 0)
    
    doc.add_paragraph(f"Datum vytvoření: {datetime.datetime.now().strftime('%d.%m.%Y')}")
    doc.add_paragraph("Dodavatel: Instalatérský a topenářský servis (Doplnit)")
    
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Light Shading Accent 1'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Název'
    hdr_cells[1].text = 'Množství'
    hdr_cells[2].text = 'Cena/ks'
    hdr_cells[3].text = 'Cena celkem (bez DPH)'
    hdr_cells[4].text = 'DPH (21%)'
    hdr_cells[5].text = 'Cena s DPH'
    
    total_bez_dph = 0
    total_s_dph = 0
    
    for item in items:
        row_cells = table.add_row().cells
        
        qty = item['quantity']
        price = item['price']
        total = qty * price
        dph = total * 0.21
        s_dph = total + dph
        
        total_bez_dph += total
        total_s_dph += s_dph
        
        row_cells[0].text = item['name']
        row_cells[1].text = str(qty)
        row_cells[2].text = f"{price:.2f} Kč"
        row_cells[3].text = f"{total:.2f} Kč"
        row_cells[4].text = f"{dph:.2f} Kč"
        row_cells[5].text = f"{s_dph:.2f} Kč"

    doc.add_paragraph("\n")
    p = doc.add_paragraph()
    p.add_run(f"Celkem bez DPH: {total_bez_dph:.2f} Kč\n").bold = True
    p.add_run(f"Celkem s DPH: {total_s_dph:.2f} Kč").bold = True
    
    filepath = f"exports/{project_name}_nabidka.docx"
    doc.save(filepath)
    return filepath

def generate_excel(project_name: str, items: list) -> str:
    wb = Workbook()
    ws = wb.active
    ws.title = "Kalkulace"
    
    headers = ["Název položky", "Obchod", "Množství", "Cena za kus", "Sazba DPH", "Celkem bez DPH", "Celkem s DPH"]
    ws.append(headers)
    
    # Formátování hlavičky
    for cell in ws[1]:
        cell.font = Font(bold=True)
        cell.fill = PatternFill(start_color="D9D9D9", end_color="D9D9D9", fill_type="solid")
    
    row_idx = 2
    for item in items:
        ws.cell(row=row_idx, column=1, value=item['name'])
        ws.cell(row=row_idx, column=2, value=item['shop'])
        ws.cell(row=row_idx, column=3, value=item['quantity'])
        ws.cell(row=row_idx, column=4, value=item['price']).number_format = '#,##0.00 "Kč"'
        ws.cell(row=row_idx, column=5, value=0.21).number_format = '0%'
        
        # Automatické vzorce pro přepočet ceny bez DPH (Množství * Cena za kus)
        ws.cell(row=row_idx, column=6, value=f"=C{row_idx}*D{row_idx}").number_format = '#,##0.00 "Kč"'
        # Vzorec pro cenu s DPH
        ws.cell(row=row_idx, column=7, value=f"=F{row_idx}*(1+E{row_idx})").number_format = '#,##0.00 "Kč"'
        
        row_idx += 1

    # Řádek celkových součtů
    ws.cell(row=row_idx, column=5, value="CELKEM:")
    ws.cell(row=row_idx, column=5).font = Font(bold=True)
    
    ws.cell(row=row_idx, column=6, value=f"=SUM(F2:F{row_idx-1})").number_format = '#,##0.00 "Kč"'
    ws.cell(row=row_idx, column=6).font = Font(bold=True)
    
    ws.cell(row=row_idx, column=7, value=f"=SUM(G2:G{row_idx-1})").number_format = '#,##0.00 "Kč"'
    ws.cell(row=row_idx, column=7).font = Font(bold=True)

    # Přizpůsobení šířky sloupců
    for col in ws.columns:
        max_length = 0
        column = col[0].column_letter 
        for cell in col:
            try:
                if len(str(cell.value)) > max_length:
                    max_length = len(str(cell.value))
            except:
                pass
        ws.column_dimensions[column].width = max_length + 2

    filepath = f"exports/{project_name}_kalkulace.xlsx"
    wb.save(filepath)
    return filepath